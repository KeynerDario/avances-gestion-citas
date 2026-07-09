import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCREENSHOTS_DIR = os.path.join(os.path.dirname(__file__), 'docs', 'screenshots')

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)
    return h

def add_image(doc, filename, caption=None, width=5.5):
    path = os.path.join(SCREENSHOTS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    else:
        doc.add_paragraph(f'[Imagen no encontrada: {filename}]')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def create_user_manual():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Portada
    title = doc.add_heading('Manual de Usuario', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)
    subtitle = doc.add_paragraph('Sistema de Gestión de Citas de Bienestar SENA')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    doc.add_paragraph('')
    ver = doc.add_paragraph('Versión 1.0 — Julio 2026')
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ver.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    # Captura de portada
    add_image(doc, '22-admin-dashboard.png', 'Vista general del Panel de Administración')
    doc.add_page_break()

    # ========================================
    # 1. INTRODUCCIÓN
    # ========================================
    add_heading(doc, '1. Introducción')
    doc.add_paragraph(
        'El Sistema de Gestión de Citas de Bienestar SENA es una plataforma web que permite '
        'a los aprendices del SENA agendar citas con profesionales de bienestar (psicología, '
        'enfermería, nutrición, etc.) y facilita la gestión, seguimiento y reporte de estas citas.'
    )
    add_heading(doc, 'Roles del Sistema', level=2)
    add_table(doc, ['Rol', 'Descripción'], [
        ['Aprendiz', 'Usuario estándar que agenda y gestiona sus citas de bienestar'],
        ['Profesional', 'Profesional de salud que atiende citas y registra notas clínicas'],
        ['Coordinación', 'Coordinador de dependencia que visualiza métricas y reportes'],
        ['Super Admin', 'Administrador del sistema con acceso total'],
    ])

    # ========================================
    # 2. ACCESO AL SISTEMA
    # ========================================
    doc.add_page_break()
    add_heading(doc, '2. Acceso al Sistema')

    add_heading(doc, '2.1 Iniciar Sesión', level=2)
    doc.add_paragraph('Abra su navegador y acceda a la URL del sistema. Verá la pantalla de inicio de sesión:')
    add_image(doc, '01-login.png', 'Figura 1. Pantalla de inicio de sesión')

    doc.add_paragraph('')
    steps = [
        'Ingrese su correo electrónico en el campo correspondiente',
        'Ingrese su contraseña',
        'Haga clic en "Iniciar Sesión"',
        'Si los datos son correctos, será redirigido a su dashboard según su rol',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    add_heading(doc, '2.2 Formulario Lleno', level=2)
    add_image(doc, '21-login-filled.png', 'Figura 2. Formulario de login con credenciales')

    add_heading(doc, '2.3 Validación de Formulario', level=2)
    doc.add_paragraph('Si intenta iniciar sesión con campos vacíos, el sistema mostrará errores:')
    add_image(doc, '06-login-validation.png', 'Figura 3. Validación de campos vacíos')

    add_heading(doc, '2.4 Funcionalidades del Login', level=2)
    for b in [
        'Mostrar/ocultar contraseña: Haga clic en el ícono del ojo',
        '¿Olvidó su contraseña? Haga clic en el enlace para recibir un correo',
        'Crear cuenta: Haga clic en "Regístrate aquí"',
    ]:
        doc.add_paragraph(b, style='List Bullet')

    # ========================================
    # 3. REGISTRO
    # ========================================
    doc.add_page_break()
    add_heading(doc, '3. Registro de Usuarios')

    add_heading(doc, '3.1 Crear una Cuenta', level=2)
    doc.add_paragraph('Haga clic en "Regístrate aquí" para ver el formulario:')
    add_image(doc, '02-register.png', 'Figura 4. Formulario de registro')

    doc.add_paragraph('')
    doc.add_paragraph('Complete los campos:')
    for f in [
        'Nombre completo: Sus nombre y apellidos',
        'Número de documento: Su cédula o documento',
        'Correo electrónico: Un correo válido',
        'Contraseña: Mínimo 6 caracteres',
        'Confirmar contraseña: Debe coincidir',
    ]:
        doc.add_paragraph(f, style='List Bullet')

    add_heading(doc, '3.2 Selección de Rol', level=2)
    doc.add_paragraph('Seleccione su rol: Aprendiz, Profesional o Coordinación.')
    add_image(doc, '07-register-profesional.png', 'Figura 5. Registro como PROFESIONAL con dependencia')

    doc.add_paragraph('')
    note = doc.add_paragraph()
    run = note.add_run('Nota: ')
    run.font.bold = True
    note.add_run('Si selecciona Profesional o Coordinación, esperará aprobación del administrador.')

    add_heading(doc, '3.3 Fortaleza de Contraseña', level=2)
    add_image(doc, '08-register-weak-password.png', 'Figura 6. Contraseña débil')
    add_image(doc, '09-register-strong-password.png', 'Figura 7. Contraseña fuerte')
    add_table(doc, ['Nivel', 'Color', 'Criterio'], [
        ['Débil', 'Rojo', '< 4 caracteres'],
        ['Regular', 'Amarillo', '4-5 caracteres'],
        ['Buena', 'Azul', '6-7 caracteres'],
        ['Fuerte', 'Verde', '8+ caracteres'],
    ])

    # ========================================
    # 4. DASHBOARD ADMIN
    # ========================================
    doc.add_page_break()
    add_heading(doc, '4. Panel de Administración')
    doc.add_paragraph('Al iniciar sesión como Super Admin, accede al panel de administración:')
    add_image(doc, '22-admin-dashboard.png', 'Figura 8. Panel de Administración — Vista General')

    add_heading(doc, '4.1 Estadísticas', level=2)
    doc.add_paragraph('La parte superior muestra tarjetas con:')
    for item in [
        'Total de usuarios registrados',
        'Usuarios activos',
        'Usuarios inactivos',
        'Nuevos en los últimos 7 días',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========================================
    # 5. GESTIÓN DE USUARIOS
    # ========================================
    doc.add_page_break()
    add_heading(doc, '5. Gestión de Usuarios')
    add_image(doc, '22-admin-dashboard.png', 'Figura 9. Pestaña de Gestión de Usuarios')

    add_heading(doc, '5.1 Filtrar por Rol', level=2)
    doc.add_paragraph('Use los botones de rol para filtrar:')
    for item in [
        'SUPERADMIN — Administradores del sistema',
        'COORDINACION — Coordinadores de dependencia',
        'PROFESIONAL — Profesionales de salud',
        'APRENDIZ — Aprendices del SENA',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, '5.2 Buscar Usuarios', level=2)
    doc.add_paragraph('Use la barra de búsqueda para encontrar usuarios por nombre, documento o email.')

    add_heading(doc, '5.3 Gestionar Usuario', level=2)
    doc.add_paragraph('Para cada usuario puede:')
    for item in [
        'Activar/desactivar con el botón de estado',
        'Expandir detalles con el chevrón',
        'Cambiar rol desde el panel expandido',
        'Asignar dependencia (solo para profesionales)',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========================================
    # 6. DEPENDENCIAS
    # ========================================
    doc.add_page_break()
    add_heading(doc, '6. Gestión de Dependencias')
    add_image(doc, '25-admin-dependencies.png', 'Figura 10. Gestión de Dependencias')
    doc.add_paragraph('Desde esta pestaña puede:')
    for item in [
        'Crear nuevas dependencias con nombre y color',
        'Editar dependencias existentes',
        'Eliminar dependencias',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========================================
    # 7. ROLES Y PERMISOS
    # ========================================
    add_heading(doc, '7. Roles y Permisos')
    add_image(doc, '26-admin-roles.png', 'Figura 11. Visualización de Roles y Permisos')
    doc.add_paragraph('Muestra los roles configurados en el sistema y sus permisos asociados.')

    # ========================================
    # 8. AUDITORÍA
    # ========================================
    doc.add_page_break()
    add_heading(doc, '8. Registro de Auditoría')
    add_image(doc, '27-admin-audit.png', 'Figura 12. Registro de Auditoría')
    doc.add_paragraph('El registro de auditoría muestra todas las acciones realizadas en el sistema:')
    for item in [
        'Crear/actualizar/eliminar usuarios',
        'Gestión de citas (crear, confirmar, cancelar)',
        'Cambios de configuración',
        'Filtrar por tipo de acción',
        'Buscar por texto',
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ========================================
    # 9. CONFIGURACIÓN
    # ========================================
    add_heading(doc, '9. Configuración del Sistema')
    add_image(doc, '28-admin-config.png', 'Figura 13. Panel de Configuración')
    doc.add_paragraph('Ajustes generales del sistema que el administrador puede modificar.')

    # ========================================
    # 10. GESTIÓN DE CITAS (APRENDIZ)
    # ========================================
    doc.add_page_break()
    add_heading(doc, '10. Gestión de Citas (Aprendiz)')

    add_heading(doc, '10.1 Solicitar una Cita', level=2)
    for i, step in enumerate([
        'Haga clic en el botón "+" o "Agendar Cita"',
        'Seleccione la dependencia',
        'Seleccione una fecha (mínimo 24h, sin fines de semana)',
        'Seleccione un horario (8:00 AM - 4:30 PM)',
        'Describe el motivo (mínimo 10 caracteres)',
        'Haga clic en "Solicitar Cita"',
    ], 1):
        doc.add_paragraph(f'{i}. {step}')

    note2 = doc.add_paragraph()
    run = note2.add_run('Importante: ')
    run.font.bold = True
    note2.add_run('Máximo 2 citas pendientes simultáneamente.')

    add_heading(doc, '10.2 Filtrar Citas', level=2)
    for f in ['Todas', 'Pendientes', 'Confirmadas', 'Completadas', 'Canceladas', 'No asistió']:
        doc.add_paragraph(f, style='List Bullet')

    # ========================================
    # 11. AGENDA DEL PROFESIONAL
    # ========================================
    add_heading(doc, '11. Agenda del Profesional')
    for a in [
        'Ver citas del día con datos del aprendiz',
        'Confirmar / Completar / Marcar inasistencia',
        'Registrar notas clínicas',
        'Consultar historial y estadísticas',
        'Configurar horarios semanales',
    ]:
        doc.add_paragraph(a, style='List Bullet')

    # ========================================
    # 12. COORDINACIÓN
    # ========================================
    add_heading(doc, '12. Panel de Coordinación')
    for c in [
        'KPIs: Total citas, pendientes, tasa cumplimiento, profesionales activos',
        'Gráficas: Citas por dependencia, tendencia mensual',
        'Tabla de rendimiento por profesional',
        'Filtros por rango de fechas',
        'Exportación a CSV',
    ]:
        doc.add_paragraph(c, style='List Bullet')

    # ========================================
    # 13. MODO OSCURO
    # ========================================
    add_heading(doc, '13. Modo Oscuro')
    doc.add_paragraph('Haga clic en el ícono sol/luna en la barra lateral para alternar el tema.')

    # ========================================
    # 14. ATAJOS DE TECLADO
    # ========================================
    add_heading(doc, '14. Atajos de Teclado')
    add_table(doc, ['Atajo', 'Acción'], [
        ['Ctrl + K', 'Abrir paleta de comandos'],
        ['Ctrl + B', 'Contraer/expandir barra lateral'],
        ['Escape', 'Cerrar modales'],
    ])

    # ========================================
    # 15. MÓVIL
    # ========================================
    doc.add_page_break()
    add_heading(doc, '15. Versión Móvil')
    doc.add_paragraph('El sistema es responsive:')
    add_image(doc, '10-mobile-login.png', 'Figura 14. Login en móvil', width=3.0)
    add_image(doc, '11-mobile-register.png', 'Figura 15. Registro en móvil', width=3.0)

    # ========================================
    # 16. SOLUCIÓN DE PROBLEMAS
    # ========================================
    add_heading(doc, '16. Solución de Problemas')
    for title, sol in [
        ('No puedo iniciar sesión', 'Verifique correo y contraseña. Use "¿Olvidaste tu contraseña?".'),
        ('La página no carga', 'Actualice (F5). Limpie caché. Verifique conexión.'),
        ('No puedo agendar cita', 'Verifique < 2 pendientes. Fecha futura. Profesional disponible.'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f'{title}: ')
        run.font.bold = True
        p.add_run(sol)

    doc.save('docs/MANUAL_USUARIO.docx')
    print('✓ Manual de usuario: docs/MANUAL_USUARIO.docx')


def create_tech_manual():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    title = doc.add_heading('Manual Técnico', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x2e)
    subtitle = doc.add_paragraph('Sistema de Gestión de Citas de Bienestar SENA')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
    doc.add_paragraph('')
    ver = doc.add_paragraph('Versión 1.0 — Julio 2026')
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ver.runs:
        run.font.size = Pt(10)
        run.font.italic = True
    doc.add_page_break()

    # 1
    add_heading(doc, '1. Arquitectura del Sistema')
    doc.add_paragraph('Arquitectura de componentes React con patrón Repository. Backend en Supabase.')
    flow = ['Autenticación via Supabase Auth', 'Perfil con rol y dependencia', 'React Router redirige según rol', 'Repository consultan Supabase', 'Realtime via WebSocket']
    for i, f in enumerate(flow, 1):
        doc.add_paragraph(f'{i}. {f}')

    # 2
    add_heading(doc, '2. Stack Tecnológico')
    add_table(doc, ['Capa', 'Tecnología', 'Versión'], [
        ['Framework', 'React', '19.2.4'], ['Build', 'Vite', '8.0.4'],
        ['Routing', 'React Router DOM', '7.14.0'], ['Forms', 'React Hook Form + Zod', '7.72.1 + 4.3.6'],
        ['Icons', 'Lucide React', '1.8.0'], ['Charts', 'Recharts', '3.8.1'],
        ['Backend', 'Supabase', '2.103.0'], ['Testing', 'Vitest', '4.1.8'],
    ])

    # 3
    doc.add_page_break()
    add_heading(doc, '3. Estructura del Proyecto')
    p = doc.add_paragraph()
    run = p.add_run('src/\n├── features/    # admin, appointments, auth, dashboard\n├── shared/      # components, styles, utils\n├── providers/   # Auth, Theme\n├── routes/      # AppRoutes, ProtectedRoute\n└── lib/         # supabase.js')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    add_heading(doc, 'Módulos', level=2)
    add_table(doc, ['Módulo', 'Repository', 'Responsabilidad'], [
        ['appointments', 'AppointmentRepository', 'CRUD citas, batch enrichment'],
        ['admin', 'AdminRepository', 'Usuarios, dependencias, auditoría'],
        ['dashboard', 'DashboardRepository', 'KPIs, métricas, CSV'],
        ['professional', 'ProfessionalRepository', 'Horarios, notas, agenda'],
    ])

    # 4
    add_heading(doc, '4. Configuración')
    add_heading(doc, 'Variables de Entorno', level=2)
    p = doc.add_paragraph()
    run = p.add_run('VITE_SUPABASE_URL=https://...\nVITE_SUPABASE_ANON_KEY=...')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    add_heading(doc, 'Comandos', level=2)
    add_table(doc, ['Comando', 'Descripción'], [
        ['npm install', 'Instalar dependencias'], ['npm run dev', 'Servidor desarrollo'],
        ['npm run test:run', 'Ejecutar tests'], ['npm run build', 'Build producción'],
        ['npm run lint', 'Verificar lint'],
    ])

    # 5
    doc.add_page_break()
    add_heading(doc, '5. Base de Datos')
    add_table(doc, ['Tabla', 'Descripción'], [
        ['profiles', 'Perfiles de usuario'], ['roles', 'Roles del sistema'],
        ['dependencies', 'Dependencias'], ['appointments', 'Citas'],
        ['professional_schedules', 'Horarios'], ['clinical_notes', 'Notas clínicas'],
        ['audit_logs', 'Auditoría'], ['system_config', 'Configuración'],
    ])

    add_heading(doc, 'Relaciones', level=2)
    p = doc.add_paragraph()
    run = p.add_run('profiles → roles (role_id)\nprofiles → dependencies (dependency_id)\nappointments → profiles (user_id, professional_id)\nappointments → dependencies (dependency_id)')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    # 6
    add_heading(doc, '6. Autenticación')
    add_table(doc, ['Ruta', 'Roles'], [
        ['/dashboard', 'APRENDIZ'], ['/professional', 'PROFESIONAL'],
        ['/coordination', 'COORDINACION, SUPERADMIN'], ['/admin', 'SUPERADMIN'],
    ])

    # 7
    add_heading(doc, '7. Patrón Repository')
    doc.add_paragraph('Cada módulo usa Repository para aislar consultas Supabase. Enrichment batch para evitar N+1.')

    # 8
    add_heading(doc, '8. Optimización')
    add_table(doc, ['Técnica', 'Descripción'], [
        ['Request Queue', 'MAX_CONCURRENT=5 para evitar ERR_INSUFFICIENT_RESOURCES'],
        ['Batch Enrichment', '3 queries batch en vez de 3N'],
        ['Lazy Loading', 'React.lazy() por ruta'],
        ['Debounce Realtime', '1s delay en eventos Supabase'],
    ])

    # 9
    add_heading(doc, '9. Testing — 108 tests')
    add_table(doc, ['Módulo', 'Tests'], [
        ['appointments', '40'], ['admin', '8'], ['dashboard', '11'],
        ['shared', '31'], ['providers', '5'], ['utils', '7'], ['auth', '6'],
    ])

    # 10
    add_heading(doc, '10. Build')
    add_table(doc, ['Archivo', 'Tamaño', 'Gzipped'], [
        ['index.js', '~187 KB', '~59 KB'], ['supabase.js', '~223 KB', '~58 KB'],
        ['index.css', '~133 KB', '~20 KB'],
    ])

    # 11
    doc.add_page_break()
    add_heading(doc, '11. Guía de Desarrollo')
    add_heading(doc, 'Nueva Feature', level=2)
    for i, s in enumerate([
        'Crear src/features/[nombre]/', 'api/[nombre].repository.js',
        'components/', 'hooks/use[Nombre].js', 'pages/[Nombre]Page.jsx',
        'Agregar ruta en AppRoutes.jsx', 'Proteger en ProtectedRoute.jsx', 'Crear tests',
    ], 1):
        doc.add_paragraph(f'{i}. {s}')

    add_heading(doc, 'Commits', level=2)
    add_table(doc, ['Tipo', 'Ejemplo'], [
        ['feat', 'feat(appointments): add cancel button'],
        ['fix', 'fix(auth): handle expired token'],
        ['test', 'test(admin): add user tests'],
    ])

    # 12
    add_heading(doc, '12. Solución de Problemas')
    for t, s in [
        ('ERR_INSUFFICIENT_RESOURCES', 'Cola de requests limita conexiones'),
        ('Lock not released', 'React Strict Mode en desarrollo'),
        ('act() warning', 'Envolver en waitFor()'),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(f'{t}: ')
        run.font.bold = True
        p.add_run(s)

    doc.save('docs/MANUAL_TECNICO.docx')
    print('✓ Manual técnico: docs/MANUAL_TECNICO.docx')


if __name__ == '__main__':
    create_user_manual()
    create_tech_manual()
    print('\n✓ Manuales generados con screenshots')
